# vectors.py - Fixed Document Processing & Search Index Manager

import os
import re
import logging
import hashlib
import time
import uuid
import threading
import tempfile
import json
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from functools import wraps
from concurrent.futures import ThreadPoolExecutor, as_completed

# Document processing imports
import PyPDF2
import docx
from docx import Document as DocxDocument
from tempfile import NamedTemporaryFile
import fitz 

# Azure imports
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.models import VectorizedQuery
from azure.search.documents.indexes.models import (
    SearchIndex,
    SearchField,
    SearchFieldDataType,
    SimpleField,
    SearchableField,
    VectorSearch,
    HnswAlgorithmConfiguration,
    VectorSearchProfile,
    SemanticConfiguration,
    SemanticSearch,
    SemanticPrioritizedFields,
    SemanticField
)
from azure.core.credentials import AzureKeyCredential
from openai import AzureOpenAI

# Langchain imports
try:
    from langchain_community.document_loaders import PyPDFLoader, UnstructuredPDFLoader, TextLoader, Docx2txtLoader
    PDF_LOADERS_AVAILABLE = True
except ImportError:
    try:
        from langchain.document_loaders import PyPDFLoader, UnstructuredPDFLoader, TextLoader, Docx2txtLoader
        PDF_LOADERS_AVAILABLE = True
    except ImportError:
        PDF_LOADERS_AVAILABLE = False
        logging.warning("PDF loaders not available. Install with: pip install langchain pypdf unstructured")

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    from langchain.text_splitter import RecursiveCharacterTextSplitter

try:
    from langchain_openai import AzureOpenAIEmbeddings
except ImportError:
    logging.error("AzureOpenAIEmbeddings not available")

from langchain.schema import Document
from dotenv import load_dotenv

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
load_dotenv()

# Production limits and security
MAX_DOCUMENT_SIZE = 50 * 1024 * 1024  # 50MB
MAX_CHUNKS_PER_DOCUMENT = 1500
MAX_PROCESSING_TIME = 300  # 5 minutes
ALLOWED_FILE_TYPES = {'.pdf', '.txt', '.docx', '.pptx'}
MAX_CONCURRENT_PROCESSING = 5

# Thread-safe processing lock
processing_lock = threading.Lock()
active_sessions = {}

def retry_on_failure(max_retries=3, delay=1):
    """Decorator for retrying failed operations"""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            for attempt in range(max_retries):
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    if attempt == max_retries - 1:
                        raise e
                    logger.warning(f"Attempt {attempt + 1} failed: {e}. Retrying in {delay}s...")
                    time.sleep(delay * (attempt + 1))
            return None
        return wrapper
    return decorator

def validate_file_security(file_path: str) -> Tuple[bool, str]:
    """Enhanced security validation for uploaded files"""
    try:
        path = Path(file_path)
        
        if not path.exists():
            return False, "File does not exist"
        
        file_size = path.stat().st_size
        if file_size > MAX_DOCUMENT_SIZE:
            return False, f"File too large: {file_size / (1024*1024):.1f}MB > {MAX_DOCUMENT_SIZE / (1024*1024)}MB"
        
        if path.suffix.lower() not in ALLOWED_FILE_TYPES:
            return False, f"File type not allowed: {path.suffix}"
        
        # Check for suspicious patterns
        suspicious_patterns = ['../', '..\\', '<script', 'javascript:', 'data:', 'vbscript:']
        if any(pattern in path.name.lower() for pattern in suspicious_patterns):
            return False, "Suspicious file name detected"
        
        if file_size < 50:
            return False, "File too small (likely empty or corrupted)"
        
        return True, "File validation passed"
        
    except Exception as e:
        return False, f"Validation error: {str(e)}"

def escape_odata_string(value: str) -> str:
    """Enhanced OData string escaping"""
    if value is None:
        return ""
    # Replace single quotes with double single quotes for OData escaping
    # Also handle other special characters
    escaped = value.replace("'", "''")
    escaped = escaped.replace('"', '""')
    escaped = escaped.replace('\\', '\\\\')
    return escaped

class SecurityError(Exception):
    """Security-related error"""
    pass

class AccuracyVectorManager:
    """Enhanced Accuracy-First Document Processing & Search Index Manager"""
    
    def __init__(
        self,
        azure_search_endpoint: str = None,
        azure_search_key: str = None,
        azure_openai_endpoint: str = None,
        azure_openai_key: str = None,
        azure_openai_api_version: str = None,
        azure_embeddings_deployment: str = None,
        index_name: str = None,
        chunk_size: int = 1200,
        chunk_overlap: int = 300,
        max_chunks: int = MAX_CHUNKS_PER_DOCUMENT
    ):
        """Initialize accuracy-focused document and vector services"""
        
        # Generate unique session ID
        self.session_id = str(uuid.uuid4())[:8]
        
        # Azure Search configuration
        self.azure_search_endpoint = azure_search_endpoint or os.getenv('AZURE_SEARCH_ENDPOINT')
        self.azure_search_key = azure_search_key or os.getenv('AZURE_SEARCH_KEY')
        
        # Azure OpenAI configuration
        self.azure_openai_endpoint = azure_openai_endpoint or os.getenv('AZURE_OPENAI_ENDPOINT')
        self.azure_openai_key = azure_openai_key or os.getenv('AZURE_OPENAI_KEY')
        self.azure_openai_api_version = azure_openai_api_version or os.getenv('AZURE_API_VERSION', '2024-12-01-preview')
        self.azure_embeddings_deployment = azure_embeddings_deployment or os.getenv('AZURE_EMBEDDINGS_DEPLOYMENT_NAME', 'text-embedding-3-small')
        
        # Validate required configuration
        if not all([self.azure_search_endpoint, self.azure_search_key, 
                   self.azure_openai_endpoint, self.azure_openai_key]):
            raise ValueError("Missing required Azure configuration. Check environment variables.")
        
        # Create unique index name for session isolation
        if index_name:
            self.index_name = self._sanitize_index_name(index_name)
        else:
            unique_id = f"{self.session_id}_{int(time.time())}"
            self.index_name = f"accuracy-index-{unique_id}"[:63]
            
        logger.info(f"AccuracyVectorManager initialized with index: {self.index_name}")
        
        # Processing parameters
        self.chunk_size = max(500, min(chunk_size, 2000))
        self.chunk_overlap = max(50, min(chunk_overlap, self.chunk_size // 2))
        self.max_chunks = max_chunks
        
        # Processing statistics
        self.stats_lock = threading.Lock()
        self.stats = {
            'documents_processed': 0,
            'total_chunks_created': 0,
            'processing_time': 0,
            'errors_encountered': 0,
            'session_id': self.session_id,
            'index_name': self.index_name
        }
        
        # Track session
        with processing_lock:
            active_sessions[self.session_id] = {
                'start_time': time.time(),
                'index_name': self.index_name,
                'stats': self.stats
            }
        
        self._initialize_components()

    def _initialize_components(self) -> None:
        """Initialize all components with retry logic"""
        self._initialize_azure_openai_client()
        self._initialize_azure_search_clients()
        self._initialize_text_splitter()
        self._initialize_embedding_model()
        self._create_search_index()

    def _initialize_azure_openai_client(self) -> None:
        """Initialize Azure OpenAI client for embeddings"""
        try:
            self.azure_openai_client = AzureOpenAI(
                api_key=self.azure_openai_key,
                api_version=self.azure_openai_api_version,
                azure_endpoint=self.azure_openai_endpoint
            )
            
            # Test embeddings
            test_response = self.azure_openai_client.embeddings.create(
                input="test",
                model=self.azure_embeddings_deployment
            )
            
            # Get embedding dimension
            self.embedding_dimension = len(test_response.data[0].embedding)
            logger.info(f"Azure OpenAI embeddings initialized: {self.embedding_dimension} dimensions")
            
        except Exception as e:
            logger.error(f"Failed to initialize Azure OpenAI client: {e}")
            raise ConnectionError(f"Azure OpenAI initialization failed: {e}")

    def _initialize_azure_search_clients(self) -> None:
        """Initialize Azure AI Search clients"""
        try:
            credential = AzureKeyCredential(self.azure_search_key)
            
            self.search_index_client = SearchIndexClient(
                endpoint=self.azure_search_endpoint,
                credential=credential
            )
            
            self.search_client = SearchClient(
                endpoint=self.azure_search_endpoint,
                index_name=self.index_name,
                credential=credential
            )
            
            logger.info(f"Azure AI Search clients initialized for index: {self.index_name}")
            
        except Exception as e:
            logger.error(f"Failed to initialize Azure Search clients: {e}")
            raise ConnectionError(f"Azure Search initialization failed: {e}")

    def _initialize_text_splitter(self) -> None:
        """Initialize text splitter with enhanced separators"""
        self.text_splitter = RecursiveCharacterTextSplitter(
            separators=[
                "\n\n\n", "\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " ", ""
            ],
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            keep_separator=True,
            add_start_index=True
        )
        logger.info(f"Text splitter initialized: {self.chunk_size} chars, {self.chunk_overlap} overlap")

    def _initialize_embedding_model(self) -> None:
        """Initialize the embedding model for LangChain integration"""
        try:
            self.embedding_model = AzureOpenAIEmbeddings(
                deployment=self.azure_embeddings_deployment,
                model="text-embedding-3-small",
                chunk_size=1,
                api_key=self.azure_openai_key,
                azure_endpoint=self.azure_openai_endpoint,
                api_version=self.azure_openai_api_version
            )
            logger.info("LangChain embedding model initialized")
        except Exception as e:
            logger.error(f"Failed to initialize embedding model: {e}")
            self.embedding_model = None

    def _sanitize_index_name(self, name: str) -> str:
        """Sanitize index name for Azure AI Search"""
        if not name or not isinstance(name, str):
            return f"accuracy-index-{self.session_id}-{int(time.time())}"[:63]
        
        # Azure AI Search naming rules: lowercase alphanumeric and hyphens only
        sanitized = re.sub(r'[^a-z0-9\-]', '-', name.lower())
        sanitized = re.sub(r'-+', '-', sanitized)  # Remove multiple consecutive hyphens
        sanitized = sanitized.strip('-')  # Remove leading/trailing hyphens
        
        # Ensure uniqueness
        if not any(char.isdigit() for char in sanitized):
            sanitized = f"{sanitized}-{int(time.time())}"
        
        return sanitized[:63]

    def _create_search_index(self) -> None:
        """Create Azure AI Search index with evaluation storage capabilities"""
        try:
            # Check if index already exists
            try:
                existing_index = self.search_index_client.get_index(self.index_name)
                logger.info(f"Using existing search index: {self.index_name}")
                return
            except Exception:
                pass  # Index doesn't exist, create new one
            
            # Define index fields for resume documents AND evaluation results
            fields = [
                SimpleField(name="id", type=SearchFieldDataType.String, key=True, sortable=True),
                SearchableField(name="content", type=SearchFieldDataType.String, searchable=True),
                SearchableField(name="title", type=SearchFieldDataType.String, searchable=True, filterable=True),
                SimpleField(name="document_type", type=SearchFieldDataType.String, filterable=True),  # "resume" or "evaluation"
                SimpleField(name="candidate_name", type=SearchFieldDataType.String, filterable=True, facetable=True),
                SimpleField(name="jd_hash", type=SearchFieldDataType.String, filterable=True),
                SimpleField(name="session_id", type=SearchFieldDataType.String, filterable=True),
                SimpleField(name="processing_timestamp", type=SearchFieldDataType.Double, sortable=True),
                
                # Evaluation-specific fields
                SimpleField(name="total_score", type=SearchFieldDataType.Double, sortable=True),
                SimpleField(name="skills_score", type=SearchFieldDataType.Double, filterable=True),
                SimpleField(name="experience_score", type=SearchFieldDataType.Double, filterable=True),
                SimpleField(name="project_score", type=SearchFieldDataType.Double, filterable=True),
                SearchableField(name="evaluation_details", type=SearchFieldDataType.String, searchable=True),
                SearchableField(name="matched_skills", type=SearchFieldDataType.String, searchable=True),
                SearchableField(name="skill_evidence", type=SearchFieldDataType.String, searchable=True),
                
                # Vector search for resume content
                SearchField(
                    name="content_vector",
                    type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
                    searchable=True,
                    vector_search_dimensions=self.embedding_dimension,
                    vector_search_profile_name="myHnswProfile"
                )
            ]
            
            # Configure vector search
            vector_search = VectorSearch(
                algorithms=[
                    HnswAlgorithmConfiguration(name="myHnsw")
                ],
                profiles=[
                    VectorSearchProfile(
                        name="myHnswProfile",
                        algorithm_configuration_name="myHnsw"
                    )
                ]
            )
            
            # Configure semantic search
            semantic_config = SemanticConfiguration(
                name="my-semantic-config",
                prioritized_fields=SemanticPrioritizedFields(
                    title_field=SemanticField(field_name="title"),
                    content_fields=[SemanticField(field_name="content")]
                )
            )
            
            semantic_search = SemanticSearch(configurations=[semantic_config])
            
            # Create index
            index = SearchIndex(
                name=self.index_name,
                fields=fields,
                vector_search=vector_search,
                semantic_search=semantic_search
            )
            
            self.search_index_client.create_index(index)
            logger.info(f"Created accuracy-focused search index: {self.index_name}")
            
        except Exception as e:
            logger.error(f"Failed to create search index {self.index_name}: {e}")
            raise ConnectionError(f"Search index creation failed: {e}")

    # ==================== ENHANCED DOCUMENT PROCESSING METHODS ====================
    
    def extract_text_from_file(self, file) -> str:
        """Enhanced text extraction from PDF or DOCX files with multiple fallback methods"""
        if not file or not hasattr(file, 'name'):
            raise ValueError("Invalid file object")
            
        ext = os.path.splitext(file.name)[-1].lower()
        text = ""

        try:
            if ext == ".pdf":
                text = self._extract_pdf_text_enhanced(file)
            elif ext == ".docx":
                text = self._extract_docx_text_enhanced(file)
            else:
                raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX are supported.")
            
            # Validate extracted text
            if not text or len(text.strip()) < 100:
                raise ValueError(f"Insufficient text extracted from {file.name}: {len(text)} characters")
                
            # Clean and normalize text
            text = self._clean_and_validate_text(text)
            
            logger.info(f"Successfully extracted {len(text)} characters from {file.name}")
            return text
            
        except Exception as e:
            logger.error(f"Text extraction failed for {file.name}: {e}")
            raise ValueError(f"Failed to extract text from {file.name}: {e}")

    def _extract_pdf_text_enhanced(self, file) -> str:
        """Enhanced PDF text extraction with multiple methods"""
        text = ""
        file.seek(0)
        
        # Method 1: Try PyMuPDF first (better text extraction)
        try:
            file_content = file.read()
            file.seek(0)
            
            # Use PyMuPDF for better text extraction
            doc = fitz.open(stream=file_content, filetype="pdf")
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                if page_text.strip():
                    text += page_text + "\n"
            doc.close()
            
            if len(text.strip()) > 100:
                return text
                
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")
        
        # Method 2: Fallback to PyPDF2
        try:
            file.seek(0)
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += page_text + "\n"
                    
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {e}")
        
        if not text.strip():
            raise ValueError("Could not extract text from PDF using any method")
            
        return text

    def _extract_docx_text_enhanced(self, file) -> str:
        """Enhanced DOCX text extraction"""
        text = ""
        
        with NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            file.seek(0)
            temp_file.write(file.read())
            temp_file_path = temp_file.name

        try:
            doc = docx.Document(temp_file_path)
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            # Extract headers/footers if available
            for section in doc.sections:
                if hasattr(section, 'header') and section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            text += para.text + "\n"
                            
        except Exception as e:
            logger.error(f"Enhanced DOCX extraction failed: {e}")
            raise
        finally:
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)

        if not text.strip():
            raise ValueError("Could not extract text from DOCX file")
            
        return text

    def process_multiple_documents(self, files) -> List[Document]:
        """Enhanced processing of multiple PDF and DOCX files with better error handling"""
        documents_with_metadata = []
        processing_errors = []

        for file in files:
            if not file or not hasattr(file, 'name'):
                processing_errors.append("Invalid file object")
                continue
                
            suffix = os.path.splitext(file.name)[-1].lower()

            if suffix not in ['.pdf', '.docx']:
                logger.warning(f"Skipping unsupported file format: {file.name}")
                processing_errors.append(f"Unsupported format: {file.name}")
                continue

            try:
                # Extract text using enhanced method
                extracted_text = self.extract_text_from_file(file)
                
                if not extracted_text or len(extracted_text.strip()) < 100:
                    processing_errors.append(f"Insufficient text in {file.name}")
                    continue

                # Create candidate name from filename
                candidate_name = file.name.rsplit('.', 1)[0]
                candidate_name = re.sub(r'[^a-zA-Z0-9\s\-_]', '', candidate_name).strip()
                
                if not candidate_name:
                    candidate_name = f"Unknown_Candidate_{int(time.time())}"

                # Create document with enhanced metadata
                document = Document(
                    page_content=extracted_text,
                    metadata={
                        'candidate_name': candidate_name,
                        'original_filename': file.name,
                        'file_type': suffix,
                        'text_length': len(extracted_text),
                        'word_count': len(extracted_text.split()),
                        'processing_timestamp': time.time(),
                        'full_resume_text': extracted_text  # Store complete text
                    }
                )

                documents_with_metadata.append(document)
                logger.info(f"Successfully processed {file.name} -> {candidate_name}")
                
            except Exception as e:
                error_msg = f"Failed to process {file.name}: {str(e)}"
                logger.error(error_msg)
                processing_errors.append(error_msg)
                continue

        if processing_errors:
            logger.warning(f"Processing errors: {processing_errors}")
            
        if not documents_with_metadata:
            raise ValueError(f"No documents could be processed. Errors: {processing_errors}")

        logger.info(f"Successfully processed {len(documents_with_metadata)} documents")
        return documents_with_metadata

    # ==================== ENHANCED SEARCH INDEX STORAGE METHODS ====================

    def get_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Enhanced embeddings generation with better error handling"""
        if not texts:
            return []
            
        try:
            # Clean texts before embedding
            cleaned_texts = []
            for text in texts:
                if not text or not text.strip():
                    cleaned_texts.append("No content available")
                else:
                    # Truncate very long texts to avoid API limits
                    cleaned_text = text.strip()[:8000]  # OpenAI limit is 8191 tokens
                    cleaned_texts.append(cleaned_text)
            
            # Process in batches to handle API limits
            batch_size = 16
            all_embeddings = []
            
            for i in range(0, len(cleaned_texts), batch_size):
                batch = cleaned_texts[i:i + batch_size]
                
                retry_count = 0
                max_retries = 3
                
                while retry_count < max_retries:
                    try:
                        response = self.azure_openai_client.embeddings.create(
                            input=batch,
                            model=self.azure_embeddings_deployment
                        )
                        
                        batch_embeddings = [data.embedding for data in response.data]
                        all_embeddings.extend(batch_embeddings)
                        break
                        
                    except Exception as e:
                        retry_count += 1
                        if retry_count >= max_retries:
                            logger.error(f"Failed to get embeddings after {max_retries} retries: {e}")
                            # Return zero vectors as fallback
                            zero_vector = [0.0] * self.embedding_dimension
                            all_embeddings.extend([zero_vector] * len(batch))
                        else:
                            wait_time = 2 ** retry_count
                            logger.warning(f"Embedding API error, retrying in {wait_time}s: {e}")
                            time.sleep(wait_time)
                
                # Small delay to respect rate limits
                time.sleep(0.1)
            
            return all_embeddings
            
        except Exception as e:
            logger.error(f"Failed to get embeddings: {e}")
            # Return zero vectors as fallback
            zero_vector = [0.0] * self.embedding_dimension
            return [zero_vector] * len(texts)

    def store_jd_requirements(self, jd_text: str, requirements: Dict) -> str:
        """Enhanced JD requirements storage with validation"""
        try:
            if not jd_text or not jd_text.strip():
                raise ValueError("JD text cannot be empty")
                
            jd_hash = hashlib.md5(jd_text.strip().encode()).hexdigest()
            doc_id = f"jd_requirements_{jd_hash}_{self.session_id}"
            
            # Validate requirements structure
            if not isinstance(requirements, dict):
                raise ValueError("Requirements must be a dictionary")
            
            # Create document for JD requirements
            document = {
                "id": doc_id,
                "content": jd_text[:5000],  # Truncate for storage
                "title": f"JD Requirements {jd_hash[:8]}",
                "document_type": "jd_requirements",
                "candidate_name": "JD_REQUIREMENTS",
                "jd_hash": jd_hash,
                "session_id": self.session_id,
                "processing_timestamp": time.time(),
                "evaluation_details": json.dumps(requirements, ensure_ascii=False),
                "content_vector": [0.0] * self.embedding_dimension  # Placeholder vector
            }
            
            result = self.search_client.upload_documents([document])
            
            if result and len(result) > 0 and result[0].succeeded:
                logger.info(f"Stored JD requirements in search index: {jd_hash}")
                return jd_hash
            else:
                raise Exception("Failed to store JD requirements in search index")
                
        except Exception as e:
            logger.error(f"Error storing JD requirements: {e}")
            raise ValueError(f"Failed to store JD requirements: {e}")

    def get_cached_jd_requirements(self, jd_text: str) -> Optional[Dict]:
        """Enhanced JD requirements retrieval with better error handling"""
        try:
            if not jd_text or not jd_text.strip():
                return None
                
            jd_hash = hashlib.md5(jd_text.strip().encode()).hexdigest()
            escaped_session_id = escape_odata_string(self.session_id)
            
            results = self.search_client.search(
                search_text="*",
                filter=f"document_type eq 'jd_requirements' and jd_hash eq '{jd_hash}' and session_id eq '{escaped_session_id}'",
                select=["evaluation_details"],
                top=1
            )
            
            for result in results:
                requirements_json = result.get("evaluation_details")
                if requirements_json:
                    try:
                        requirements = json.loads(requirements_json)
                        logger.info(f"Retrieved cached JD requirements: {jd_hash}")
                        return requirements
                    except json.JSONDecodeError as e:
                        logger.error(f"Failed to parse cached JD requirements: {e}")
                        return None
                        
        except Exception as e:
            logger.error(f"Error retrieving cached JD requirements: {e}")
            
        return None

    def store_candidate_evaluation(
        self, 
        candidate_name: str, 
        jd_hash: str, 
        evaluation_result: Dict,
        resume_text: str = ""
    ) -> None:
        """Enhanced candidate evaluation storage with validation"""
        try:
            if not candidate_name or not candidate_name.strip():
                raise ValueError("Candidate name cannot be empty")
                
            if not jd_hash:
                raise ValueError("JD hash cannot be empty")
                
            if not isinstance(evaluation_result, dict):
                raise ValueError("Evaluation result must be a dictionary")
            
            # Sanitize candidate_name: replace invalid characters with underscores
            safe_candidate_name = re.sub(r'[^a-zA-Z0-9_\-]', '_', candidate_name.strip())
            doc_id = f"evaluation_{jd_hash}_{safe_candidate_name}_{self.session_id}_{int(time.time())}"
            
            # Create document for evaluation results
            document = {
                "id": doc_id,
                "content": resume_text[:1000] if resume_text else "",  # Truncated for search
                "title": f"Evaluation: {candidate_name}",
                "document_type": "evaluation",
                "candidate_name": candidate_name,
                "jd_hash": jd_hash,
                "session_id": self.session_id,
                "processing_timestamp": time.time(),
                "total_score": float(evaluation_result.get("total_score", 0)),
                "skills_score": float(evaluation_result.get("skills_score", 0)),
                "experience_score": float(evaluation_result.get("experience_score", 0)),
                "project_score": float(evaluation_result.get("project_score", 0)),
                "evaluation_details": json.dumps(evaluation_result, ensure_ascii=False),
                "matched_skills": json.dumps(evaluation_result.get("matched_skills", []), ensure_ascii=False),
                "skill_evidence": json.dumps(evaluation_result.get("skill_evidence", {}), ensure_ascii=False),
                "content_vector": [0.0] * self.embedding_dimension  # Placeholder vector
            }
            
            result = self.search_client.upload_documents([document])
            
            if result and len(result) > 0 and result[0].succeeded:
                logger.info(f"Stored evaluation for {candidate_name} in search index")
            else:
                logger.warning(f"Failed to store evaluation for {candidate_name}")
                
        except Exception as e:
            logger.error(f"Error storing candidate evaluation for {candidate_name}: {e}")

    def get_cached_evaluation(self, candidate_name: str, jd_hash: str) -> Optional[Dict]:
        """Enhanced candidate evaluation retrieval with proper escaping"""
        try:
            if not candidate_name or not jd_hash:
                return None
                
            escaped_candidate_name = escape_odata_string(candidate_name)
            escaped_session_id = escape_odata_string(self.session_id)
            
            results = self.search_client.search(
                search_text="*",
                filter=f"document_type eq 'evaluation' and candidate_name eq '{escaped_candidate_name}' and jd_hash eq '{jd_hash}' and session_id eq '{escaped_session_id}'",
                select=["evaluation_details", "processing_timestamp"],
                top=1,
                order_by=["processing_timestamp desc"]
            )
            
            for result in results:
                evaluation_json = result.get("evaluation_details")
                if evaluation_json:
                    try:
                        evaluation = json.loads(evaluation_json)
                        logger.info(f"Retrieved cached evaluation for {candidate_name}")
                        return evaluation
                    except json.JSONDecodeError as e:
                        logger.error(f"Failed to parse cached evaluation: {e}")
                        return None
                        
        except Exception as e:
            logger.error(f"Error retrieving cached evaluation for {candidate_name}: {e}")
            
        return None

    @retry_on_failure(max_retries=2)
    def create_embeddings_from_files(self, files: List) -> str:
        """Enhanced embeddings creation with comprehensive error handling"""
        start_time = time.time()
        
        try:
            if not files:
                raise ValueError("No files provided for processing")
            
            # Process multiple documents with enhanced error handling
            documents_with_metadata = self.process_multiple_documents(files)
            
            if not documents_with_metadata:
                raise ValueError("No valid documents found in uploaded files")
                
            # Split into chunks
            splits = self.text_splitter.split_documents(documents_with_metadata)
            
            # Apply limits
            if len(splits) > self.max_chunks:
                logger.warning(f"Too many chunks ({len(splits)}), limiting to {self.max_chunks}")
                step = len(splits) // self.max_chunks
                splits = splits[::max(1, step)][:self.max_chunks]
            
            # Clean and validate splits
            valid_splits = []
            for i, split in enumerate(splits):
                if not split.page_content or len(split.page_content.strip()) < 150:
                    continue
                
                clean_content = self._clean_and_validate_text(split.page_content)
                if not clean_content or len(clean_content.strip()) < 100:
                    continue
                
                split.page_content = clean_content
                split.metadata.update({
                    "chunk_id": i,
                    "chunk_size": len(clean_content),
                    "word_count": len(clean_content.split()),
                    "processing_timestamp": time.time(),
                    "session_id": self.session_id,
                    "index_name": self.index_name
                })
                
                valid_splits.append(split)
            
            if not valid_splits:
                raise ValueError("No valid chunks created from documents")
            
            logger.info(f"Created {len(valid_splits)} chunks for indexing")
            
            # Generate embeddings with enhanced error handling
            texts = [split.page_content for split in valid_splits]
            embeddings = self.get_embeddings(texts)
            
            if len(embeddings) != len(texts):
                logger.warning(f"Embedding count mismatch: {len(embeddings)} vs {len(texts)}")
                # Pad with zero vectors if needed
                while len(embeddings) < len(texts):
                    embeddings.append([0.0] * self.embedding_dimension)
            
            # Prepare documents for indexing
            documents = []
            for i, (split, embedding) in enumerate(zip(valid_splits, embeddings)):
                candidate_name = split.metadata.get('candidate_name', 'unknown')
                # Sanitize candidate_name: replace invalid characters with underscores
                safe_candidate_name = re.sub(r'[^a-zA-Z0-9_\-]', '_', candidate_name)
                doc_id = f"resume_{safe_candidate_name}_{i}_{int(time.time())}"
                
                document = {
                    "id": doc_id,
                    "content": split.page_content,
                    "title": split.metadata.get('candidate_name', 'Unknown'),
                    "document_type": "resume",
                    "candidate_name": split.metadata.get('candidate_name', 'Unknown'),
                    "session_id": self.session_id,
                    "processing_timestamp": split.metadata.get('processing_timestamp', time.time()),
                    "content_vector": embedding
                }
                documents.append(document)
            
            # Upload to Azure AI Search with retry
            result = self.search_client.upload_documents(documents)
            
            # Check for errors
            failed_uploads = [r for r in result if not r.succeeded]
            if failed_uploads:
                logger.warning(f"Some documents failed to upload: {len(failed_uploads)}")
                for failed in failed_uploads:
                    logger.error(f"Upload failed: {failed.error_message if hasattr(failed, 'error_message') else 'Unknown error'}")
            
            successful_uploads = len([r for r in result if r.succeeded])
            logger.info(f"Successfully indexed {successful_uploads} resume documents")
            
            # Update statistics
            processing_time = time.time() - start_time
            with self.stats_lock:
                self.stats.update({
                    'documents_processed': self.stats['documents_processed'] + len(files),
                    'total_chunks_created': self.stats['total_chunks_created'] + len(valid_splits),
                    'processing_time': self.stats['processing_time'] + processing_time
                })
            
            return f"Successfully processed {len(files)} documents into {successful_uploads} searchable chunks"
            
        except Exception as e:
            with self.stats_lock:
                self.stats['errors_encountered'] += 1
            logger.error(f"Embedding creation failed: {e}")
            raise Exception(f"Document processing failed: {e}")

    def get_candidate_documents(self, candidate_name: str) -> List[Document]:
        """Enhanced candidate document retrieval with better error handling"""
        try:
            if not candidate_name or not candidate_name.strip():
                return []
                
            # Escape candidate name for Azure Search filter
            escaped_candidate_name = escape_odata_string(candidate_name)
            escaped_session_id = escape_odata_string(self.session_id)
            
            documents = self.search_client.search(
                search_text="*",
                filter=f"session_id eq '{escaped_session_id}' and document_type eq 'resume' and candidate_name eq '{escaped_candidate_name}'",
                select=["content", "candidate_name", "processing_timestamp"],
                order_by=["processing_timestamp desc"]
            )
            
            candidate_documents = []
            for doc in documents:
                if doc.get("content"):
                    candidate_documents.append(Document(
                        page_content=doc["content"],
                        metadata={
                            "candidate_name": doc.get("candidate_name", candidate_name),
                            "processing_timestamp": doc.get("processing_timestamp", time.time())
                        }
                    ))
            
            logger.info(f"Retrieved {len(candidate_documents)} documents for {candidate_name}")
            return candidate_documents
            
        except Exception as e:
            logger.error(f"Failed to retrieve candidate documents for {candidate_name}: {e}")
            return []
    
    def _clean_and_validate_text(self, text: str) -> str:
        """Enhanced text cleaning and validation"""
        if not text or not text.strip():
            return ""
        
        try:
            # Remove security threats
            text = re.sub(r'<script[^>]*>.*?</script>', '[SCRIPT_REMOVED]', text, flags=re.IGNORECASE | re.DOTALL)
            text = re.sub(r'javascript:', '[JS_REMOVED]', text, flags=re.IGNORECASE)
            text = re.sub(r'data:(?!image)', '[DATA_REMOVED]', text, flags=re.IGNORECASE)
            
            # Basic PII masking
            text = re.sub(r'\b\d{3}[-.\s]?\d{2}[-.\s]?\d{4}\b', '[SSN_REDACTED]', text)
            text = re.sub(r'\b\d{4}[-.\s]?\d{4}[-.\s]?\d{4}[-.\s]?\d{4}\b', '[CARD_REDACTED]', text)
            
            # Clean excessive whitespace
            text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
            text = re.sub(r'[ \t]+', ' ', text)
            text = re.sub(r'\r\n', '\n', text)
            
            # Remove special characters that might cause parsing issues
            text = re.sub(r'[^\w\s\-.,;:()[]{}/@#$%^&*+=|\\<>?!"\'`~]', ' ', text)
            
            # Normalize encoding
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            
            return text.strip()
            
        except Exception as e:
            logger.warning(f"Text cleaning failed: {e}")
            return text.strip() if text else ""

    def get_stats(self) -> Dict[str, Any]:
        """Enhanced statistics retrieval"""
        with self.stats_lock:
            stats_copy = self.stats.copy()
        
        # Add session info
        session_info = active_sessions.get(self.session_id, {})
        
        return {
            **stats_copy,
            "index_name": self.index_name,
            "max_chunks_limit": self.max_chunks,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "session_active_time": time.time() - session_info.get('start_time', time.time()),
            "embedding_dimension": getattr(self, 'embedding_dimension', 'unknown'),
            "session_start_time": session_info.get('start_time', time.time())
        }

    def cleanup_session(self) -> None:
        """Enhanced session cleanup with better error handling"""
        try:
            with processing_lock:
                if self.session_id in active_sessions:
                    del active_sessions[self.session_id]
            
            # DELETE the entire index
            try:
                logger.info(f"Deleting index: {self.index_name}")
                self.search_index_client.delete_index(self.index_name)
                logger.info(f"Successfully deleted index: {self.index_name}")
            except Exception as e:
                logger.warning(f"Index deletion failed for {self.index_name}: {e}")
            
            # Close clients
            try:
                if hasattr(self, 'search_client'):
                    self.search_client.close()
                if hasattr(self, 'search_index_client'):
                    self.search_index_client.close()
            except Exception as e:
                logger.warning(f"Client cleanup failed: {e}")
            
        except Exception as e:
            logger.error(f"Session cleanup failed: {e}")

def get_active_sessions_info() -> Dict[str, Any]:
    """Get active sessions information"""
    with processing_lock:
        return {
            "total_active_sessions": len(active_sessions),
            "sessions": {
                session_id: {
                    "duration": time.time() - info["start_time"],
                    "index_name": info["index_name"],
                    "documents_processed": info["stats"]["documents_processed"]
                }
                for session_id, info in active_sessions.items()
            }
        }